from __future__ import annotations

import re
import os
import subprocess
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

import mistune
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.shared import Cm, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree
from mathml2omml import convert as mathml_to_omml
from mistune.plugins import plugin_table

from word_format_config import DEFAULT_WORD_FORMAT, WordFormatConfig


MATH_BLOCK_PATTERN = re.compile(r"\$\$\s*(.+?)\s*\$\$", re.DOTALL)
MATH_BLOCK_PLACEHOLDER_PATTERN = re.compile(r"@@MB(\d+)@@")
INLINE_MATH_PLACEHOLDER_PATTERN = re.compile(r"@@MI(\d+)@@")
H1_PATTERN = re.compile(r"(?m)^#\s+(.+)\s*$")
TAG_PATTERN = re.compile(r"\\tag\{([^}]+)\}")
CODE_LINE_NO_PATTERN = re.compile(r"^\s*\d+→(?: (?=\S))?")

_RECOVERY_XML_PARSER = etree.XMLParser(recover=True)


@dataclass(frozen=True)
class PreparedMarkdown:
    text: str
    math_blocks: list[str]
    inline_math: list[str]

@dataclass(frozen=True)
class TocEntry:
    level: int
    label: str
    bookmark_name: str


@dataclass
class RenderContext:
    toc_entries: list[TocEntry]
    toc_pos: int = 0
    next_bookmark_id: int = 1
    seen_first_h1: bool = False
    last_was_page_break: bool = False
    in_references: bool = False
    last_was_blank_line: bool = False
    table_counter: int = 0
    pending_table_no: Optional[int] = None
    table_number_map: dict[int, int] = field(default_factory=dict)
    last_token_was_table: bool = False
    base_path: Path = field(default_factory=lambda: Path("."))
    config: WordFormatConfig = field(default_factory=lambda: DEFAULT_WORD_FORMAT)


def _add_blank_line(doc: Document, ctx: Optional[RenderContext] = None) -> None:
    if ctx and ctx.last_was_blank_line:
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if ctx:
        ctx.last_was_blank_line = True


def _get_or_add_rfonts(rpr) -> OxmlElement:
    for child in rpr:
        if child.tag == qn("w:rFonts"):
            return child
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
    return rfonts


def _apply_run_fonts(run, *, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = _get_or_add_rfonts(rpr)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _parse_xml_with_recovery(xml: str):
    try:
        return parse_xml(xml)
    except Exception:
        recovered = etree.fromstring(xml.encode("utf-8"), parser=_RECOVERY_XML_PARSER)
        recovered_xml = etree.tostring(recovered, encoding="unicode")
        return parse_xml(recovered_xml)


def _extract_math_blocks(markdown: str) -> tuple[str, list[str]]:
    blocks: list[str] = []

    def _repl(match: re.Match[str]) -> str:
        blocks.append(match.group(1))
        return f"\n\n@@MB{len(blocks) - 1}@@\n\n"

    return MATH_BLOCK_PATTERN.sub(_repl, markdown), blocks


def _extract_inline_math(markdown: str) -> tuple[str, list[str]]:
    out: list[str] = []
    inline: list[str] = []

    in_code_block = False

    for line in markdown.splitlines(True):
        stripped = line.lstrip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            out.append(line)
            continue
        if in_code_block:
            out.append(line)
            continue

        i = 0
        n = len(line)
        in_codespan = False
        while i < n:
            ch = line[i]

            if ch == "`":
                in_codespan = not in_codespan
                out.append(ch)
                i += 1
                continue

            if in_codespan:
                out.append(ch)
                i += 1
                continue

            if ch == "$" and (i == 0 or line[i - 1] != "\\"):
                if i + 1 < n and line[i + 1] == "$":
                    out.append("$")
                    i += 1
                    continue
                j = i + 1
                while j < n:
                    if line[j] == "$" and line[j - 1] != "\\":
                        break
                    j += 1
                if j < n:
                    latex = line[i + 1 : j]
                    inline.append(latex)
                    out.append(f"@@MI{len(inline) - 1}@@")
                    i = j + 1
                    continue

            out.append(ch)
            i += 1

    return "".join(out), inline


def prepare_markdown(markdown: str) -> PreparedMarkdown:
    prepared, math_blocks = _extract_math_blocks(markdown)
    prepared, inline_math = _extract_inline_math(prepared)
    return PreparedMarkdown(text=prepared, math_blocks=math_blocks, inline_math=inline_math)

def _normalize_latex(latex: str) -> tuple[str, Optional[str]]:
    s = latex.strip()
    tag = None
    m = TAG_PATTERN.search(s)
    if m:
        tag = m.group(1).strip()
        s = TAG_PATTERN.sub("", s).strip()
    s = (
        s.replace("tok,pos", "tok_pos")
        .replace("tok,neg", "tok_neg")
        .replace("seq,pos", "seq_pos")
        .replace("seq,neg", "seq_neg")
    )
    return s, tag


def _latex_to_mathml(latex: str, *, display: bool) -> str:
    try:
        script = (
            "const katex=require('katex');"
            "const fs=require('fs');"
            "const tex=fs.readFileSync(0,'utf8');"
            "const display=process.env.KATEX_DISPLAY==='1';"
            "const out=katex.renderToString(tex,{output:'mathml',displayMode:display,throwOnError:false,strict:'ignore'});"
            "process.stdout.write(out);"
        )
        result = subprocess.run(
            ["node", "-e", script],
            input=latex,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            env={**os.environ, "KATEX_DISPLAY": "1" if display else "0"},
            cwd=str(Path(__file__).resolve().parent),
        )
        if result.returncode == 0 and "<math" in result.stdout and "</math>" in result.stdout:
            start = result.stdout.find("<math")
            end = result.stdout.rfind("</math>")
            raw_mathml = result.stdout[start : end + len("</math>")]
            try:
                root = etree.fromstring(raw_mathml.encode("utf-8"))
                ns = {"m": "http://www.w3.org/1998/Math/MathML"}

                for anno in root.xpath(".//*[local-name()='annotation' or local-name()='annotation-xml']"):
                    parent = anno.getparent()
                    if parent is not None:
                        parent.remove(anno)

                semantics = root.xpath("./m:semantics", namespaces=ns)
                if semantics:
                    sem = semantics[0]
                    kept = None
                    for child in list(sem):
                        if etree.QName(child).localname not in ("annotation", "annotation-xml"):
                            kept = child
                            break
                    root.clear()
                    if kept is not None:
                        root.append(kept)

                for mspace in root.xpath(".//*[local-name()='mspace' and @linebreak]"):
                    parent = mspace.getparent()
                    if parent is not None:
                        parent.remove(mspace)

                mml_ns = "http://www.w3.org/1998/Math/MathML"

                for mover in root.xpath(".//*[local-name()='mover' and (@accent='true' or @accent='1')]"):
                    kids = list(mover)
                    if len(kids) >= 2 and etree.QName(kids[1]).localname == "mo":
                        if (kids[1].text or "").strip() == "^":
                            kids[1].text = "ˆ"

                for elem in root.xpath(".//*[local-name()='msub']"):
                    kids = list(elem)
                    if len(kids) > 2:
                        base = kids[0]
                        subkids = kids[1:]
                        for k in subkids:
                            elem.remove(k)
                        mrow = etree.Element(f"{{{mml_ns}}}mrow")
                        for k in subkids:
                            mrow.append(k)
                        elem.append(mrow)

                for elem in root.xpath(".//*[local-name()='msup']"):
                    kids = list(elem)
                    if len(kids) > 2:
                        base = kids[0]
                        supkids = kids[1:]
                        for k in supkids:
                            elem.remove(k)
                        mrow = etree.Element(f"{{{mml_ns}}}mrow")
                        for k in supkids:
                            mrow.append(k)
                        elem.append(mrow)

                for elem in root.xpath(".//*[local-name()='msubsup']"):
                    kids = list(elem)
                    if len(kids) > 3:
                        base = kids[0]
                        sub = kids[1]
                        extras = kids[3:]
                        for k in extras:
                            elem.remove(k)
                        mrow = etree.Element(f"{{{mml_ns}}}mrow")
                        for k in extras:
                            mrow.append(k)
                        elem.append(mrow)

                return etree.tostring(root, encoding="unicode")
            except Exception:
                return raw_mathml
    except Exception:
        pass
    return latex_to_mathml(latex)


def _append_math_inline(paragraph, latex: str) -> None:
    normalized, _ = _normalize_latex(latex)
    mathml = _latex_to_mathml(normalized, display=False)
    omml = mathml_to_omml(mathml).strip()
    omml = omml.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
    element = _parse_xml_with_recovery(omml)
    run = paragraph.add_run()
    run._r.append(element)


def _append_math_block(paragraph, latex: str) -> None:
    normalized, _ = _normalize_latex(latex)
    mathml = _latex_to_mathml(normalized, display=True)
    omml = mathml_to_omml(mathml).strip()
    math_para = f"<m:oMathPara {nsdecls('m')}>{omml}</m:oMathPara>"
    element = _parse_xml_with_recovery(math_para)
    paragraph._p.append(element)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_equation_with_number(doc: Document, latex: str, tag: str, *, config: WordFormatConfig) -> None:
    section = doc.sections[-1]
    usable_emu = section.page_width - section.left_margin - section.right_margin - Cm(0.2)
    right_cm = float(usable_emu) / 360000.0
    center_cm = right_cm / 2.0

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(center_cm), alignment=WD_TAB_ALIGNMENT.CENTER, leader=WD_TAB_LEADER.SPACES)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(right_cm), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES)

    _apply_run_fonts(p.add_run("\t"), east_asia=config.body_east_asia_font, ascii_font=config.body_ascii_font)
    _append_math_inline(p, latex)
    _apply_run_fonts(p.add_run("\t"), east_asia=config.body_east_asia_font, ascii_font=config.body_ascii_font)
    num_run = p.add_run(f"({tag})")
    num_run.bold = False
    num_run.italic = False
    _apply_run_fonts(num_run, east_asia=config.body_east_asia_font, ascii_font=config.body_ascii_font)


def _iter_text_with_placeholders(text: str) -> Iterable[tuple[str, str]]:
    cursor = 0
    for match in INLINE_MATH_PLACEHOLDER_PATTERN.finditer(text):
        start, end = match.span()
        if start > cursor:
            yield ("text", text[cursor:start])
        yield ("inline_math", match.group(1))
        cursor = end
    if cursor < len(text):
        yield ("text", text[cursor:])


def _render_text(
    paragraph,
    text: str,
    prepared: PreparedMarkdown,
    *,
    bold: bool = False,
    italic: bool = False,
    allow_italic: bool = True,
    east_asia: str = "宋体",
    enable_citation_superscript: bool = True,
    table_number_map: Optional[dict[int, int]] = None,
) -> None:
    citation_pattern = re.compile(r"\[(\s*\d+(?:\s*[-–]\s*\d+)?(?:\s*,\s*\d+(?:\s*[-–]\s*\d+)?)*)\]")
    table_ref_pattern = re.compile(r"表\s*(\d+)")

    for kind, value in _iter_text_with_placeholders(text):
        if kind == "text":
            if not value:
                continue
            if table_number_map:
                def _table_repl(m: re.Match[str]) -> str:
                    old = int(m.group(1))
                    new = table_number_map.get(old)
                    return f"表{new}" if new is not None else m.group(0)
                value = table_ref_pattern.sub(_table_repl, value)
            if not enable_citation_superscript:
                run = paragraph.add_run(value)
                run.bold = bold or None
                run.italic = (italic and allow_italic) or None
                _apply_run_fonts(run, east_asia=east_asia)
                continue
            cursor = 0
            for match in citation_pattern.finditer(value):
                start, end = match.span()
                if start > cursor:
                    seg = value[cursor:start]
                    if seg:
                        run = paragraph.add_run(seg)
                        run.bold = bold or None
                        run.italic = (italic and allow_italic) or None
                        _apply_run_fonts(run, east_asia=east_asia)

                cite = match.group(1).strip()
                cite_run = paragraph.add_run(f"［{cite}］")
                cite_run.bold = bold or None
                cite_run.italic = (italic and allow_italic) or None
                cite_run.font.superscript = True
                _apply_run_fonts(cite_run, east_asia=east_asia)
                cursor = end

            if cursor < len(value):
                seg = value[cursor:]
                if seg:
                    run = paragraph.add_run(seg)
                    run.bold = bold or None
                    run.italic = (italic and allow_italic) or None
                    _apply_run_fonts(run, east_asia=east_asia)
        else:
            idx = int(value)
            if 0 <= idx < len(prepared.inline_math):
                _append_math_inline(paragraph, prepared.inline_math[idx])


def _render_inlines(
    paragraph,
    inlines: list[dict],
    prepared: PreparedMarkdown,
    *,
    bold: bool = False,
    italic: bool = False,
    allow_italic: bool = True,
    east_asia: str = "宋体",
    enable_citation_superscript: bool = True,
    table_number_map: Optional[dict[int, int]] = None,
    base_path: Path = Path("."),
) -> None:
    for node in inlines:
        node_type = node.get("type")
        if node_type == "text":
            _render_text(
                paragraph,
                node.get("text", ""),
                prepared,
                bold=bold,
                italic=italic,
                allow_italic=allow_italic,
                east_asia=east_asia,
                enable_citation_superscript=enable_citation_superscript,
                table_number_map=table_number_map,
            )
        elif node_type == "strong":
            _render_inlines(
                paragraph,
                node.get("children", []),
                prepared,
                bold=True or bold,
                italic=italic,
                allow_italic=allow_italic,
                east_asia=east_asia,
                enable_citation_superscript=enable_citation_superscript,
                table_number_map=table_number_map,
            )
        elif node_type == "emphasis":
            _render_inlines(
                paragraph,
                node.get("children", []),
                prepared,
                bold=bold,
                italic=(True or italic) if allow_italic else False,
                allow_italic=allow_italic,
                east_asia=east_asia,
                enable_citation_superscript=enable_citation_superscript,
                table_number_map=table_number_map,
            )
        elif node_type == "codespan":
            run = paragraph.add_run(node.get("text", ""))
            run.bold = bold or None
            run.italic = (italic and allow_italic) or None
            _apply_run_fonts(run, east_asia=east_asia)
        elif node_type == "link":
            _render_inlines(
                paragraph,
                node.get("children", []),
                prepared,
                bold=bold,
                italic=italic,
                allow_italic=allow_italic,
                east_asia=east_asia,
                enable_citation_superscript=enable_citation_superscript,
                table_number_map=table_number_map,
                base_path=base_path,
            )
        elif node_type == "image":
            src = node.get("src") or ""
            if src.startswith("http://") or src.startswith("https://"):
                run = paragraph.add_run(f"[{src}]")
                run.bold = bold or None
                run.italic = (italic and allow_italic) or None
                _apply_run_fonts(run, east_asia=east_asia)
                continue
            img_path = Path(src)
            if not img_path.is_absolute():
                img_path = (base_path / img_path).resolve()
            if not img_path.exists():
                raise FileNotFoundError(str(img_path))
            doc = paragraph.part.document
            section = doc.sections[-1]
            max_width = section.page_width - section.left_margin - section.right_margin
            run = paragraph.add_run()
            run.add_picture(str(img_path), width=max_width)
        elif node_type == "linebreak":
            paragraph.add_run().add_break()
        else:
            if "children" in node:
                _render_inlines(
                    paragraph,
                    node.get("children", []),
                    prepared,
                    bold=bold,
                    italic=italic,
                    allow_italic=allow_italic,
                    east_asia=east_asia,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                    base_path=base_path,
                )
            elif "text" in node:
                _render_text(
                    paragraph,
                    str(node.get("text", "")),
                    prepared,
                    bold=bold,
                    italic=italic,
                    allow_italic=allow_italic,
                    east_asia=east_asia,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                )


def _extract_plain_text(inlines: list[dict], prepared: PreparedMarkdown) -> str:
    out: list[str] = []
    for node in inlines:
        t = node.get("type")
        if t == "text":
            text = node.get("text", "")
            text = INLINE_MATH_PLACEHOLDER_PATTERN.sub("", text)
            out.append(text)
        elif "children" in node:
            out.append(_extract_plain_text(node.get("children", []), prepared))
        elif "text" in node:
            out.append(str(node.get("text", "")))
    return "".join(out)

def _extract_raw_text(inlines: list[dict]) -> str:
    out: list[str] = []
    for node in inlines:
        t = node.get("type")
        if "text" in node:
            out.append(node.get("text", ""))
        if "children" in node and isinstance(node.get("children"), list):
            out.append(_extract_raw_text(node.get("children", [])))
        elif t == "linebreak":
            out.append("\n")
    return "".join(out)


def _preserve_leading_spaces(text: str) -> str:
    if not text:
        return ""
    i = 0
    n = len(text)
    while i < n and text[i] == " ":
        i += 1
    if i == 0:
        return text
    return ("\u00A0" * i) + text[i:]


def _normalize_code_block_text(text: str) -> str:
    lines = text.splitlines()
    cleaned: list[str] = []
    for line in lines:
        cleaned.append(CODE_LINE_NO_PATTERN.sub("", line))
    return "\n".join(cleaned)


def _set_cell_shading(cell, *, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_horizontal_margins_twips(cell, *, left: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    def _set_side(name: str, value: int) -> None:
        el = tc_mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(int(value)))
        el.set(qn("w:type"), "dxa")

    _set_side("left", left)
    _set_side("right", right)


def _set_table_code_box(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)

    tbl_borders = OxmlElement("w:tblBorders")

    def _border(tag: str, *, val: str, sz: str, color: str) -> OxmlElement:
        b = OxmlElement(tag)
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        return b

    tbl_borders.append(_border("w:top", val="single", sz="6", color="BFBFBF"))
    tbl_borders.append(_border("w:left", val="single", sz="6", color="BFBFBF"))
    tbl_borders.append(_border("w:bottom", val="single", sz="6", color="BFBFBF"))
    tbl_borders.append(_border("w:right", val="single", sz="6", color="BFBFBF"))
    tbl_borders.append(_border("w:insideH", val="none", sz="0", color="auto"))
    tbl_borders.append(_border("w:insideV", val="single", sz="6", color="E0E0E0"))

    tbl_pr.append(tbl_borders)


_PY_KEYWORDS = {
    "False",
    "None",
    "True",
    "and",
    "as",
    "assert",
    "async",
    "await",
    "break",
    "class",
    "continue",
    "def",
    "del",
    "elif",
    "else",
    "except",
    "finally",
    "for",
    "from",
    "global",
    "if",
    "import",
    "in",
    "is",
    "lambda",
    "nonlocal",
    "not",
    "or",
    "pass",
    "raise",
    "return",
    "try",
    "while",
    "with",
    "yield",
}

_BASH_KEYWORDS = {
    "case",
    "cd",
    "do",
    "done",
    "echo",
    "elif",
    "else",
    "esac",
    "export",
    "fi",
    "for",
    "function",
    "if",
    "in",
    "local",
    "return",
    "set",
    "then",
    "while",
}


def _split_unquoted_comment(line: str) -> tuple[str, str]:
    in_single = False
    in_double = False
    escaped = False
    for idx, ch in enumerate(line):
        if escaped:
            escaped = False
            continue
        if ch == "\\":
            escaped = True
            continue
        if ch == "'" and not in_double:
            in_single = not in_single
            continue
        if ch == '"' and not in_single:
            in_double = not in_double
            continue
        if ch == "#" and not in_single and not in_double:
            return line[:idx], line[idx:]
    return line, ""


def _split_leading_spaces(line: str) -> tuple[int, str]:
    i = 0
    visual = 0
    n = len(line)
    while i < n and line[i] in {" ", "\t"}:
        if line[i] == "\t":
            visual += 4
        else:
            visual += 1
        i += 1
    return visual, line[i:]


def _highlight_code_line(line: str, *, lang: str) -> list[tuple[str, RGBColor]]:
    colors = {
        "default": RGBColor(0, 0, 0),
        "comment": RGBColor(106, 153, 85),
        "string": RGBColor(206, 145, 120),
        "keyword": RGBColor(197, 134, 192),
        "number": RGBColor(181, 206, 168),
        "variable": RGBColor(156, 220, 254),
    }

    code_part, comment_part = _split_unquoted_comment(line)
    segments: list[tuple[str, RGBColor]] = []

    if lang in {"python", "py"}:
        patterns: list[tuple[re.Pattern[str], str]] = [
            (re.compile(r"('''[\s\S]*?'''|\"\"\"[\s\S]*?\"\"\"|'[^'\\]*(?:\\.[^'\\]*)*'|\"[^\"\\]*(?:\\.[^\"\\]*)*\")"), "string"),
            (re.compile(r"\b\d+(?:\.\d+)?\b"), "number"),
            (re.compile(r"\b[A-Za-z_][A-Za-z0-9_]*\b"), "ident"),
        ]
        keyword_set = _PY_KEYWORDS
    else:
        patterns = [
            (re.compile(r"('(?:\\'|[^'])*'|\"(?:\\\"|[^\"])*\")"), "string"),
            (re.compile(r"\$\{[^}]+\}|\$[A-Za-z_][A-Za-z0-9_]*"), "variable"),
            (re.compile(r"\b\d+(?:\.\d+)?\b"), "number"),
            (re.compile(r"\b[A-Za-z_][A-Za-z0-9_]*\b"), "ident"),
        ]
        keyword_set = _BASH_KEYWORDS

    i = 0
    while i < len(code_part):
        best = None
        best_kind = None
        for pat, kind in patterns:
            m = pat.search(code_part, i)
            if m is None:
                continue
            if best is None or m.start() < best.start():
                best = m
                best_kind = kind
        if best is None or best_kind is None:
            segments.append((code_part[i:], colors["default"]))
            break
        if best.start() > i:
            segments.append((code_part[i:best.start()], colors["default"]))
        token_text = best.group(0)
        if best_kind == "ident":
            segments.append((token_text, colors["keyword"] if token_text in keyword_set else colors["default"]))
        else:
            segments.append((token_text, colors[best_kind]))
        i = best.end()

    if comment_part:
        segments.append((comment_part, colors["comment"]))
    return segments


def _add_code_run(paragraph, text: str, *, color: RGBColor, config: WordFormatConfig) -> None:
    run = paragraph.add_run(text)
    run.italic = False
    run.bold = False
    run.font.size = Pt(config.code_font_size_pt)
    run.font.color.rgb = color
    _apply_run_fonts(run, east_asia=config.code_font_name, ascii_font=config.code_font_name)


def _render_code_block(doc: Document, *, code_text: str, lang: str, ctx: Optional[RenderContext]) -> None:
    config = ctx.config if ctx else DEFAULT_WORD_FORMAT
    section = doc.sections[-1]
    max_width_twips = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    gutter_width = Cm(config.code_gutter_width_cm)
    gutter_twips = int(gutter_width.twips)
    code_twips = int(max_width_twips - gutter_twips)
    code_width = Cm(code_twips / 567.0)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = gutter_width
    table.columns[1].width = code_width
    _set_table_code_box(table)
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(max_width_twips)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    col1 = OxmlElement("w:gridCol")
    col1.set(qn("w:w"), str(int(gutter_twips)))
    tbl_grid.append(col1)
    col2 = OxmlElement("w:gridCol")
    col2.set(qn("w:w"), str(int(code_twips)))
    tbl_grid.append(col2)

    def _set_cell_width_twips(cell, twips: int) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(twips))
        tc_w.set(qn("w:type"), "dxa")

    lines = code_text.splitlines()
    if not lines:
        lines = [""]

    for idx, line in enumerate(lines, start=1):
        row = table.add_row()
        row.cells[0].width = gutter_width
        row.cells[1].width = code_width
        _set_cell_width_twips(row.cells[0], gutter_twips)
        _set_cell_width_twips(row.cells[1], code_twips)

        ln_cell = row.cells[0]
        code_cell = row.cells[1]
        _set_cell_shading(ln_cell, fill="F2F2F2")
        _set_cell_shading(code_cell, fill="FAFAFA")
        _set_cell_horizontal_margins_twips(ln_cell, left=0, right=0)

        ln_p = ln_cell.paragraphs[0]
        ln_p.paragraph_format.first_line_indent = Cm(0)
        ln_p.paragraph_format.left_indent = Cm(0)
        ln_p.paragraph_format.right_indent = Cm(0)
        ln_p.paragraph_format.space_before = Pt(0)
        ln_p.paragraph_format.space_after = Pt(0)
        ln_p.paragraph_format.line_spacing = 1.0
        ln_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ln_run = ln_p.add_run(str(idx))
        ln_run.italic = False
        ln_run.bold = False
        ln_run.font.size = Pt(config.code_line_number_font_size_pt)
        ln_run.font.color.rgb = RGBColor(128, 128, 128)
        _apply_run_fonts(ln_run, east_asia=config.code_font_name, ascii_font=config.code_font_name)

        code_p = code_cell.paragraphs[0]
        code_p.paragraph_format.first_line_indent = Cm(0)
        code_p.paragraph_format.left_indent = Cm(0)
        code_p.paragraph_format.right_indent = Cm(0)
        code_p.paragraph_format.space_before = Pt(0)
        code_p.paragraph_format.space_after = Pt(0)
        code_p.paragraph_format.line_spacing = 1.0

        leading_count, rest = _split_leading_spaces(line)
        if leading_count:
            _add_code_run(code_p, "\u00A0" * leading_count, color=RGBColor(0, 0, 0), config=config)
        for seg_text, seg_color in _highlight_code_line(rest, lang=lang):
            if seg_text:
                _add_code_run(code_p, seg_text, color=seg_color, config=config)

    _add_blank_line(doc, ctx)


def _set_table_three_line(table, has_header: bool = True) -> None:
    # Set overall table borders to none
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    
    # Remove existing borders
    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)
        
    tbl_borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tbl_borders.append(border)

    # Set top border (1.5pt)
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "12")  # 12 / 8 = 1.5pt
    top_border.set(qn("w:space"), "0")
    top_border.set(qn("w:color"), "000000")
    tbl_borders.append(top_border)

    # Set bottom border (1.5pt)
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "12")
    bottom_border.set(qn("w:space"), "0")
    bottom_border.set(qn("w:color"), "000000")
    tbl_borders.append(bottom_border)
    
    tbl_pr.append(tbl_borders)

    # Set header bottom border (0.75pt) if applicable
    if has_header and len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            
            # Remove existing tcBorders
            existing_tc_borders = tc_pr.find(qn("w:tcBorders"))
            if existing_tc_borders is not None:
                tc_pr.remove(existing_tc_borders)
                
            tc_borders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")  # 6 / 8 = 0.75pt
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
            tc_borders.append(bottom)
            tc_pr.append(tc_borders)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _set_table_keep_together(table) -> None:
    rows = list(table.rows)
    for row in rows:
        _set_row_cant_split(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.keep_together = True
                pf.keep_with_next = True


def _render_table(
    doc: Document,
    token: dict,
    prepared: PreparedMarkdown,
    *,
    enable_citation_superscript: bool = True,
    table_number_map: Optional[dict[int, int]] = None,
) -> None:
    table_head = None
    table_body = None
    for child in token.get("children", []):
        if child.get("type") == "table_head":
            table_head = child
        elif child.get("type") == "table_body":
            table_body = child

    head_cells = []
    if table_head and table_head.get("children"):
        head_cells = table_head["children"]

    body_rows = []
    if table_body:
        body_rows = table_body.get("children", [])

    cols = len(head_cells)
    if cols == 0 and body_rows:
        cols = len(body_rows[0].get("children", []))
    rows = (1 if head_cells else 0) + len(body_rows)
    if rows == 0 or cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)
    row_offset = 0

    def _collect_text(nodes: list[dict]) -> str:
        buf: list[str] = []
        for nd in nodes:
            if "text" in nd:
                buf.append(nd.get("text", ""))
            if "children" in nd:
                buf.append(_collect_text(nd.get("children", [])))
        return "".join(buf)

    if head_cells:
        for c, cell_token in enumerate(head_cells):
            cell = table.cell(0, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            children = cell_token.get("children", [])
            flat = _collect_text(children)
            if "@@MI" in flat:
                _render_text(
                    p,
                    flat,
                    prepared,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                )
            else:
                _render_inlines(
                    p,
                    children,
                    prepared,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                )
            
            # Make header bold
            for run in p.runs:
                run.bold = True
        row_offset = 1

    for r, row_token in enumerate(body_rows):
        for c, cell_token in enumerate(row_token.get("children", [])):
            cell = table.cell(r + row_offset, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            children = cell_token.get("children", [])
            flat = _collect_text(children)
            if "@@MI" in flat:
                _render_text(
                    p,
                    flat,
                    prepared,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                )
            else:
                _render_inlines(
                    p,
                    children,
                    prepared,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                )
 
    _set_table_three_line(table, has_header=bool(head_cells))
    _set_table_keep_together(table)


def _render_list(
    doc: Document,
    token: dict,
    prepared: PreparedMarkdown,
    *,
    enable_citation_superscript: bool = True,
    table_number_map: Optional[dict[int, int]] = None,
    base_path: Path = Path("."),
) -> None:
    ordered = bool(token.get("ordered"))
    if ordered:
        start = int(token.get("start") or 1)
        for idx, item in enumerate(token.get("children", [])):
            p = doc.add_paragraph()
            num_run = p.add_run(f"{start + idx}. ")
            num_run.bold = False
            num_run.italic = False
            _apply_run_fonts(num_run, east_asia="宋体")
            for child in item.get("children", []):
                if child.get("type") == "block_text":
                    _render_inlines(
                        p,
                        child.get("children", []),
                        prepared,
                        enable_citation_superscript=enable_citation_superscript,
                        table_number_map=table_number_map,
                        base_path=base_path,
                    )
                else:
                    _render_blocks(doc, [child], prepared, base_path=base_path)
        return

    style = "List Bullet"
    for item in token.get("children", []):
        p = doc.add_paragraph(style=style)
        for child in item.get("children", []):
            if child.get("type") == "block_text":
                _render_inlines(
                    p,
                    child.get("children", []),
                    prepared,
                    enable_citation_superscript=enable_citation_superscript,
                    table_number_map=table_number_map,
                    base_path=base_path,
                )
            else:
                _render_blocks(doc, [child], prepared, base_path=base_path)


def _render_paragraph(
    doc: Document,
    token: dict,
    prepared: PreparedMarkdown,
    *,
    enable_citation_superscript: bool = True,
    table_number_map: Optional[dict[int, int]] = None,
    base_path: Path = Path("."),
    config: WordFormatConfig = DEFAULT_WORD_FORMAT,
) -> None:
    inlines = token.get("children", [])
    plain = _extract_plain_text(inlines, prepared).strip()

    match = MATH_BLOCK_PLACEHOLDER_PATTERN.search(plain)
    if match and plain == f"@@MB{match.group(1)}@@":
        idx = int(match.group(1))
        if 0 <= idx < len(prepared.math_blocks):
            latex = prepared.math_blocks[idx]
            _, tag = _normalize_latex(latex)
            if tag:
                _add_equation_with_number(doc, latex, tag, config=config)
            else:
                p = doc.add_paragraph()
                _append_math_block(p, latex)
        return

    p = doc.add_paragraph()
    _render_inlines(
        p,
        inlines,
        prepared,
        enable_citation_superscript=enable_citation_superscript,
        table_number_map=table_number_map,
        base_path=base_path,
    )


def _render_heading(doc: Document, token: dict) -> None:
    level = int(token.get("level", 1))
    level = max(1, min(9, level))
    p = doc.add_heading("", level=level)
    return p


def _add_bookmark(paragraph, *, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_pageref_field(paragraph, bookmark_name: str, *, east_asia: str) -> None:
    _add_field_run(paragraph, f" PAGEREF {bookmark_name} \\h ", east_asia=east_asia)


def _render_blocks(
    doc: Document,
    ast: list[dict],
    prepared: PreparedMarkdown,
    ctx: Optional[RenderContext] = None,
    *,
    base_path: Path = Path("."),
) -> None:
    base_path = (ctx.base_path if ctx else base_path).resolve()
    i = 0
    while i < len(ast):
        token = ast[i]
        t = token.get("type")

        if t == "heading":
            if ctx:
                ctx.pending_table_no = None
                ctx.last_token_was_table = False
            level = int(token.get("level", 1))
            if level <= 0:
                i += 1
                continue
            heading_text = _extract_plain_text(token.get("children", []), prepared).strip()
            if ctx and level == 1:
                normalized = re.sub(r"\s+", "", heading_text).lower()
                ctx.in_references = ("参考文献" in normalized) or (normalized == "references") or normalized.startswith("references")
                if not ctx.seen_first_h1:
                    _add_blank_line(doc, ctx)
                elif not ctx.last_was_page_break:
                    doc.add_page_break()
                    ctx.last_was_page_break = True
                    ctx.last_was_blank_line = False
                    _add_blank_line(doc, ctx)
                ctx.seen_first_h1 = True
            if ctx and level == 2:
                _add_blank_line(doc, ctx)
            heading_level = min(level, 4)
            p = doc.add_heading("", level=heading_level)
            if ctx:
                ctx.last_was_blank_line = False
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_config = ctx.config if ctx else DEFAULT_WORD_FORMAT
            _render_inlines(
                p,
                token.get("children", []),
                prepared,
                allow_italic=False,
                east_asia=heading_config.heading_east_asia_font,
                enable_citation_superscript=True,
                table_number_map=(ctx.table_number_map if ctx else None),
                base_path=base_path,
            )
            if level in (1, 2):
                _add_blank_line(doc, ctx)
            if ctx and level in (1, 2) and ctx.toc_pos < len(ctx.toc_entries):
                entry = ctx.toc_entries[ctx.toc_pos]
                if entry.level == level:
                    _add_bookmark(p, name=entry.bookmark_name, bookmark_id=ctx.next_bookmark_id)
                    ctx.next_bookmark_id += 1
                    ctx.toc_pos += 1
            i += 1
            continue

        if t == "table":
            enable = False if (ctx and ctx.in_references) else True
            table_map = ctx.table_number_map if ctx else None
            _render_table(doc, token, prepared, enable_citation_superscript=enable, table_number_map=table_map)
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = False
                ctx.table_counter += 1
                ctx.pending_table_no = ctx.table_counter
                ctx.last_token_was_table = True
            i += 1
            continue

        if t == "blank_line":
            if ctx and ctx.pending_table_no is not None:
                i += 1
                continue
            doc.add_paragraph()
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = True
                ctx.last_token_was_table = False
            i += 1
            continue

        if t == "paragraph":
            enable = False if (ctx and ctx.in_references) else True
            table_map = ctx.table_number_map if ctx else None
            if ctx and ctx.pending_table_no is not None:
                inlines = token.get("children", [])
                plain = _extract_plain_text(inlines, prepared).strip()
                if plain.startswith("表"):
                    raw = _extract_raw_text(inlines).strip()
                    m = re.match(r"^表\s*(\d+)?\s*[:：]?\s*(.*)$", raw)
                    if m:
                        rest = (m.group(2) or "").strip()
                        caption_text = f"表{ctx.pending_table_no}：{rest}" if rest else f"表{ctx.pending_table_no}："
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.paragraph_format.first_line_indent = Cm(0)
                        cap.paragraph_format.left_indent = Cm(0)
                        cap.paragraph_format.right_indent = Cm(0)
                        cap.paragraph_format.keep_together = True
                        _render_text(
                            cap,
                            caption_text,
                            prepared,
                            enable_citation_superscript=False,
                            table_number_map=None,
                        )
                        ctx.pending_table_no = None
                        _add_blank_line(doc, ctx)
                        i += 1
                        continue
                ctx.pending_table_no = None

            children = token.get("children", [])
            if len(children) == 1 and isinstance(children[0], dict) and children[0].get("type") == "image":
                img = children[0]
                src = img.get("src") or ""
                j = i + 1
                while j < len(ast) and ast[j].get("type") == "blank_line":
                    j += 1
                caption_token = ast[j] if j < len(ast) else None
                caption_plain = ""
                if caption_token and caption_token.get("type") == "paragraph":
                    caption_plain = _extract_plain_text(caption_token.get("children", []), prepared).strip()
                if caption_plain and (
                    re.match(r"^图\s*\d+\s*[:：]", caption_plain)
                    or re.match(r"^Figure\s*\d+\s*[:：]", caption_plain, flags=re.IGNORECASE)
                ):
                    img_path = Path(src)
                    if not img_path.is_absolute():
                        img_path = (base_path / img_path).resolve()
                    if not img_path.exists():
                        raise FileNotFoundError(str(img_path))

                    section = doc.sections[-1]
                    max_width = section.page_width - section.left_margin - section.right_margin

                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_p.paragraph_format.first_line_indent = Cm(0)
                    pic_p.paragraph_format.left_indent = Cm(0)
                    pic_p.paragraph_format.right_indent = Cm(0)
                    pic_p.paragraph_format.keep_together = True
                    pic_p.paragraph_format.keep_with_next = True
                    pic_p.paragraph_format.space_after = Pt(0)
                    pic_p.add_run().add_picture(str(img_path), width=max_width)

                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.first_line_indent = Cm(0)
                    cap_p.paragraph_format.left_indent = Cm(0)
                    cap_p.paragraph_format.right_indent = Cm(0)
                    cap_p.paragraph_format.keep_together = True
                    cap_p.paragraph_format.space_before = Pt(0)
                    cap_p.paragraph_format.space_after = Pt(0)
                    _render_inlines(
                        cap_p,
                        caption_token.get("children", []),
                        prepared,
                        enable_citation_superscript=enable,
                        table_number_map=table_map,
                        base_path=base_path,
                    )
                    _add_blank_line(doc, ctx)

                    i = j + 1
                    if ctx:
                        ctx.last_was_page_break = False
                        ctx.last_token_was_table = False
                    continue

            _render_paragraph(
                doc,
                token,
                prepared,
                enable_citation_superscript=enable,
                table_number_map=table_map,
                base_path=base_path,
                config=(ctx.config if ctx else DEFAULT_WORD_FORMAT),
            )
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = False
                ctx.last_token_was_table = False
            i += 1
            continue

        if t == "list":
            enable = False if (ctx and ctx.in_references) else True
            table_map = ctx.table_number_map if ctx else None
            _render_list(doc, token, prepared, enable_citation_superscript=enable, table_number_map=table_map, base_path=base_path)
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = False
                ctx.pending_table_no = None
                ctx.last_token_was_table = False
            i += 1
            continue

        if t == "block_text":
            p = doc.add_paragraph()
            enable = False if (ctx and ctx.in_references) else True
            table_map = ctx.table_number_map if ctx else None
            _render_inlines(p, token.get("children", []), prepared, enable_citation_superscript=enable, table_number_map=table_map, base_path=base_path)
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = False
                ctx.pending_table_no = None
                ctx.last_token_was_table = False
            i += 1
            continue

        if t == "thematic_break":
            doc.add_paragraph()
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = False
                ctx.pending_table_no = None
                ctx.last_token_was_table = False
            i += 1
            continue

        if t == "block_code":
            code_text = _normalize_code_block_text(str(token.get("text", "")))
            lang = str(token.get("info", "") or "").strip().lower()
            _render_code_block(doc, code_text=code_text, lang=lang, ctx=ctx)
            if ctx:
                ctx.last_was_page_break = False
                ctx.last_was_blank_line = True
                ctx.pending_table_no = None
                ctx.last_token_was_table = False
            i += 1
            continue

        children = token.get("children")
        if isinstance(children, list):
            _render_blocks(doc, children, prepared, ctx=ctx, base_path=base_path)
        i += 1


def _set_style_font(style, *, east_asia: str, ascii_font: str, size_pt: float, bold: Optional[bool] = None) -> None:
    font = style.font
    font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    font.name = ascii_font
    font.italic = False
    font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    rfonts = _get_or_add_rfonts(rpr)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _set_style_paragraph(style, *, line_spacing: float, space_before_pt: float, space_after_pt: float, first_line_indent_cm: Optional[float] = None) -> None:
    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)


def _configure_document_styles(doc: Document, *, config: WordFormatConfig) -> None:
    _set_style_font(
        doc.styles["Normal"],
        east_asia=config.body_east_asia_font,
        ascii_font=config.body_ascii_font,
        size_pt=config.body_font_size_pt,
    )
    _set_style_paragraph(
        doc.styles["Normal"],
        line_spacing=config.body_line_spacing,
        space_before_pt=0,
        space_after_pt=0,
        first_line_indent_cm=config.body_first_line_indent_cm,
    )

    if "Title" in doc.styles:
        _set_style_font(
            doc.styles["Title"],
            east_asia=config.title_east_asia_font,
            ascii_font=config.title_ascii_font,
            size_pt=config.title_font_size_pt,
            bold=True,
        )
        _set_style_paragraph(doc.styles["Title"], line_spacing=config.body_line_spacing, space_before_pt=0, space_after_pt=0, first_line_indent_cm=None)

    for name, size, east_asia, bold, after in [
        ("Heading 1", config.heading_1_size_pt, config.heading_east_asia_font, True, 0),
        ("Heading 2", config.heading_2_size_pt, config.heading_east_asia_font, True, 0),
        ("Heading 3", config.heading_3_size_pt, config.heading_east_asia_font, True, 0),
        ("Heading 4", config.heading_4_size_pt, config.heading_east_asia_font, True, 0),
    ]:
        if name in doc.styles:
            _set_style_font(doc.styles[name], east_asia=east_asia, ascii_font=config.heading_ascii_font, size_pt=size, bold=bold)
            _set_style_paragraph(doc.styles[name], line_spacing=config.body_line_spacing, space_before_pt=0, space_after_pt=after, first_line_indent_cm=None)
            doc.styles[name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if name == "Heading 1" else WD_ALIGN_PARAGRAPH.LEFT

    if "TOC Heading" in doc.styles:
        _set_style_font(
            doc.styles["TOC Heading"],
            east_asia=config.toc_heading_east_asia_font,
            ascii_font=config.toc_heading_ascii_font,
            size_pt=config.toc_heading_size_pt,
            bold=True,
        )
        _set_style_paragraph(
            doc.styles["TOC Heading"],
            line_spacing=config.body_line_spacing,
            space_before_pt=0,
            space_after_pt=config.toc_heading_space_after_pt,
            first_line_indent_cm=None,
        )

    for toc_style, east_asia, size in [
        ("TOC 1", config.toc_level1_east_asia_font, config.toc_level1_size_pt),
        ("TOC 2", config.toc_level2_east_asia_font, config.toc_level2_size_pt),
        ("TOC 3", config.toc_level3_east_asia_font, config.toc_level3_size_pt),
    ]:
        if toc_style in doc.styles:
            _set_style_font(doc.styles[toc_style], east_asia=east_asia, ascii_font=config.toc_level_font_ascii, size_pt=size, bold=False)
            _set_style_paragraph(doc.styles[toc_style], line_spacing=config.body_line_spacing, space_before_pt=0, space_after_pt=0, first_line_indent_cm=None)


def _configure_section_layout(section, *, config: WordFormatConfig) -> None:
    section.top_margin = Cm(config.margin_top_cm)
    section.bottom_margin = Cm(config.margin_bottom_cm)
    section.left_margin = Cm(config.margin_left_cm)
    section.right_margin = Cm(config.margin_right_cm)
    section.gutter = Cm(config.gutter_cm)
    section.header_distance = Cm(config.header_distance_cm)
    section.footer_distance = Cm(config.footer_distance_cm)


def _set_section_page_numbering(section, *, start: int, fmt: str) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def _set_paragraph_bottom_border(paragraph, *, config: WordFormatConfig) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(config.header_bottom_border_size))
    bottom.set(qn("w:space"), str(config.header_bottom_border_space))
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_field_run(paragraph, instr: str, *, east_asia: str = "宋体") -> None:
    run = paragraph.add_run()
    _apply_run_fonts(run, east_asia=east_asia)
    run.italic = False
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = OxmlElement("w:t")
    placeholder.text = " "
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _setup_header_footer(
    section,
    *,
    config: WordFormatConfig,
    show_page_number: bool,
    roman: bool,
    hyphen: bool,
) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    header_p.text = config.header_text
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_bottom_border(header_p, config=config)
    if header_p.runs:
        run = header_p.runs[0]
        _apply_run_fonts(run, east_asia=config.header_font_east_asia, ascii_font=config.header_font_ascii)
        run.font.size = Pt(config.header_footer_font_size_pt)

    footer_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not show_page_number:
        return

    if hyphen:
        footer_p.add_run("-")
    if roman:
        _add_field_run(footer_p, " PAGE \\* roman ")
    else:
        _add_field_run(footer_p, " PAGE ")
    if hyphen:
        footer_p.add_run("-")

    for run in footer_p.runs:
        _apply_run_fonts(run, east_asia=config.header_font_east_asia, ascii_font=config.header_font_ascii)
        run.font.size = Pt(config.header_footer_font_size_pt)


def _compute_toc_tab_pos_cm(section) -> float:
    usable_emu = section.page_width - section.left_margin - section.right_margin - Cm(0.2)
    return float(usable_emu) / 360000.0


def _add_manual_toc(doc: Document, section, toc_entries: list[TocEntry], *, config: WordFormatConfig) -> None:
    _add_blank_line(doc)
    heading = doc.add_paragraph(config.toc_title_text, style="Heading 1" if "Heading 1" in doc.styles else None)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = config.toc_title_bold
        run.italic = False
        run.font.size = Pt(config.toc_title_size_pt)
        _apply_run_fonts(run, east_asia=config.toc_title_east_asia_font, ascii_font=config.toc_title_ascii_font)
    _add_blank_line(doc)
    tab_pos_cm = _compute_toc_tab_pos_cm(section)
    for entry in toc_entries:
        toc_east_asia = config.toc_level1_east_asia_font if entry.level == 1 else config.toc_level2_east_asia_font
        style_name = "TOC 1" if entry.level == 1 else "TOC 2"
        p = doc.add_paragraph(style=style_name if style_name in doc.styles else None)
        if entry.level == 2:
            p.paragraph_format.left_indent = Cm(config.toc_level2_left_indent_cm)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos_cm), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        label_run = p.add_run(entry.label)
        label_run.italic = False
        _apply_run_fonts(label_run, east_asia=toc_east_asia, ascii_font=config.toc_level_font_ascii)

        tab_run = p.add_run("\t")
        tab_run.italic = False
        _apply_run_fonts(tab_run, east_asia=toc_east_asia, ascii_font=config.toc_level_font_ascii)
        _add_pageref_field(p, entry.bookmark_name, east_asia=toc_east_asia)


def _parse_markdown_to_ast(prepared: PreparedMarkdown) -> list[dict]:
    parser = mistune.create_markdown(renderer=mistune.AstRenderer(), plugins=[plugin_table])
    return parser(prepared.text)

def build_toc_entries(ast: list[dict], prepared: PreparedMarkdown) -> list[TocEntry]:
    entries: list[TocEntry] = []
    chapter_no = 0
    section_no = 0

    def _clean_chapter_title(s: str) -> str:
        original = s.strip()
        cleaned = re.sub(r"^第\s*\d+\s*章", "", original).strip()
        cleaned = re.sub(r"^第\d+章", "", cleaned).strip()
        return cleaned or original

    def _clean_section_title(s: str) -> str:
        original = s.strip()
        cleaned = re.sub(r"^第\s*\d+[．\.]\d+\s*节", "", original).strip()
        cleaned = re.sub(r"^\d+[．\.]\d+\s*", "", cleaned).strip()
        return cleaned or original

    def _walk(nodes: list[dict]) -> None:
        nonlocal chapter_no, section_no
        for token in nodes:
            t = token.get("type")
            if t == "heading":
                level = int(token.get("level", 1))
                if level not in (1, 2):
                    continue
                text = _extract_plain_text(token.get("children", []), prepared).strip()
                if level == 1:
                    chapter_no += 1
                    section_no = 0
                    title = _clean_chapter_title(text) or text
                    label = f"第 {chapter_no} 章 {title}"
                    entries.append(TocEntry(level=1, label=label, bookmark_name=f"toc_chap_{chapter_no}"))
                else:
                    if chapter_no == 0:
                        chapter_no = 1
                    section_no += 1
                    title = _clean_section_title(text) or text
                    label = f"第 {chapter_no}.{section_no} 节 {title}"
                    entries.append(TocEntry(level=2, label=label, bookmark_name=f"toc_sec_{chapter_no}_{section_no}"))
            else:
                children = token.get("children")
                if isinstance(children, list):
                    _walk(children)

    _walk(ast)
    return entries

def build_table_number_map(ast: list[dict], prepared: PreparedMarkdown) -> dict[int, int]:
    table_no = 0
    mapping: dict[int, int] = {}

    for i, token in enumerate(ast):
        if token.get("type") != "table":
            continue
        table_no += 1
        j = i + 1
        while j < len(ast) and ast[j].get("type") == "blank_line":
            j += 1
        if j >= len(ast) or ast[j].get("type") != "paragraph":
            continue
        inlines = ast[j].get("children", [])
        plain = _extract_plain_text(inlines, prepared).strip()
        m = re.match(r"^表\s*(\d+)", plain)
        if m:
            mapping[int(m.group(1))] = table_no

    return mapping


def _split_markdown_sections(raw: str) -> tuple[str, str, str, str, str]:
    h1s = list(H1_PATTERN.finditer(raw))
    if not h1s:
        raise ValueError("未找到一级标题（# 标题）。")

    zh_title = h1s[0].group(1).strip()

    zh_idx = None
    en_idx = None
    for idx, m in enumerate(h1s):
        header = m.group(1).strip()
        if zh_idx is None and "摘要" in header:
            zh_idx = idx
            continue
        if en_idx is None and "摘要" not in header and header.lower().startswith("abstract"):
            en_idx = idx

    if zh_idx is None:
        return zh_title, "", "", "", raw[h1s[0].end() :].lstrip()

    zh_start = h1s[zh_idx].end()
    zh_end = h1s[zh_idx + 1].start() if zh_idx + 1 < len(h1s) else len(raw)

    en_title = ""
    en_md = ""
    remaining = ""

    if en_idx is not None and en_idx > zh_idx:
        en_title_idx = en_idx - 1
        if en_title_idx > zh_idx:
            zh_end = h1s[en_title_idx].start()
            en_title = h1s[en_title_idx].group(1).strip()
        en_start = h1s[en_idx].end()
        en_end = h1s[en_idx + 1].start() if en_idx + 1 < len(h1s) else len(raw)
        en_md = raw[en_start:en_end].lstrip()
        remaining = raw[en_end:].lstrip()
    else:
        remaining = raw[zh_end:].lstrip()

    zh_md = raw[zh_start:zh_end].lstrip()
    return zh_title, zh_md, en_title, en_md, remaining


def convert_markdown_to_docx(input_path: Path, output_path: Path, *, config: WordFormatConfig = DEFAULT_WORD_FORMAT) -> Path:
    raw = input_path.read_text(encoding="utf-8", errors="replace")
    zh_title, zh_abstract_md, en_title, en_abstract_md, main_md = _split_markdown_sections(raw)
    base_path = input_path.parent.resolve()

    doc = Document()
    _configure_document_styles(doc, config=config)

    front_section = doc.sections[0]
    _configure_section_layout(front_section, config=config)
    _setup_header_footer(front_section, config=config, show_page_number=True, roman=True, hyphen=False)
    _set_section_page_numbering(front_section, start=1, fmt="roman")

    _add_blank_line(doc)
    zh_title_p = doc.add_paragraph()
    zh_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zh_run = zh_title_p.add_run(zh_title)
    zh_run.bold = True
    zh_run.italic = False
    zh_run.font.size = Pt(config.title_font_size_pt)
    _apply_run_fonts(zh_run, east_asia=config.title_east_asia_font, ascii_font=config.title_ascii_font)

    _add_blank_line(doc)
    abs_heading = doc.add_paragraph()
    abs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_run = abs_heading.add_run(config.abstract_heading_text_zh)
    abs_run.bold = True
    abs_run.italic = False
    abs_run.font.size = Pt(config.abstract_heading_size_pt)
    _apply_run_fonts(abs_run, east_asia=config.abstract_heading_font_east_asia, ascii_font=config.abstract_heading_font_ascii)
    _add_blank_line(doc)
    if zh_abstract_md.strip():
        prepared_abs = prepare_markdown(zh_abstract_md)
        ast_abs = _parse_markdown_to_ast(prepared_abs)
        _render_blocks(doc, ast_abs, prepared_abs, base_path=base_path)

    doc.add_page_break()
    if en_abstract_md.strip():
        _add_blank_line(doc)
        en_title_text = en_title.strip() if en_title.strip() else "Entropy-Focused Group Relative Policy Optimization"
        en_title_p = doc.add_paragraph()
        en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        en_title_run = en_title_p.add_run(en_title_text)
        en_title_run.bold = True
        en_title_run.italic = False
        en_title_run.font.size = Pt(config.title_font_size_pt)
        _apply_run_fonts(en_title_run, east_asia=config.title_ascii_font, ascii_font=config.title_ascii_font)

        _add_blank_line(doc)
        en_heading = doc.add_paragraph()
        en_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        en_abs_run = en_heading.add_run(config.abstract_heading_text_en)
        en_abs_run.bold = True
        en_abs_run.italic = False
        en_abs_run.font.size = Pt(config.abstract_heading_size_pt)
        _apply_run_fonts(en_abs_run, east_asia=config.abstract_heading_font_east_asia, ascii_font=config.abstract_heading_font_ascii)
        _add_blank_line(doc)
        prepared_en = prepare_markdown(en_abstract_md)
        ast_en = _parse_markdown_to_ast(prepared_en)
        _render_blocks(doc, ast_en, prepared_en, base_path=base_path)
        doc.add_page_break()
    prepared_main = prepare_markdown(main_md)
    ast_main = _parse_markdown_to_ast(prepared_main)
    table_number_map = build_table_number_map(ast_main, prepared_main)
    toc_entries = build_toc_entries(ast_main, prepared_main)
    _add_manual_toc(doc, front_section, toc_entries, config=config)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section_layout(main_section, config=config)
    _setup_header_footer(main_section, config=config, show_page_number=True, roman=False, hyphen=True)
    _set_section_page_numbering(main_section, start=1, fmt="decimal")


    ctx_main = RenderContext(toc_entries=toc_entries, table_number_map=table_number_map, config=config)
    ctx_main.base_path = base_path
    _render_blocks(doc, ast_main, prepared_main, ctx=ctx_main, base_path=base_path)

    saved_path = output_path
    try:
        doc.save(saved_path)
    except PermissionError:
        saved_path = output_path.with_name(f"{output_path.stem}_new{output_path.suffix}")
        doc.save(saved_path)

    with zipfile.ZipFile(saved_path) as zf:
        xml = zf.read("word/document.xml")
    if b"@@MI" in xml or b"@@MB" in xml:
        raise RuntimeError("仍存在未被替换的数学占位符。")
    return saved_path


if __name__ == "__main__":
    base_dir = Path(__file__).resolve().parent
    input_md = base_dir / "毕设论文草稿.md"
    output_docx = base_dir / "毕设论文草稿.docx"
    saved = convert_markdown_to_docx(input_md, output_docx, config=DEFAULT_WORD_FORMAT)
    print(str(saved))
